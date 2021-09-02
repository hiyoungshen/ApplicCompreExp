# coding=utf-8
from tkinter import *
from tkinter.filedialog import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
from openpyxl import load_workbook
import random


class GenerateQuestions():
    """
    count: 生成试卷数目
    single_choice_count: 每份试卷生成单选题数目
    multiple_choice_count： 每份试卷生成多选题数目
    judge_count: 每份试卷生成判断题数目
    """
    def __init__(self, count=10, single_choice_count=20,multiple_choice_count=20, judge_count=10) -> None:
        self.count=count
        self.single_choice_count=single_choice_count
        self.multiple_choice_count=multiple_choice_count
        self.judge_count=judge_count
        pass
    """
    questions: Document()， 生成的问题
    answers: Document(), 生成的答案
    single_choice_sheet, 题目集
    """
    def generateSingleChoose(self, questions, answers, single_choice_sheet):
        cur_answer = u''
        cur_num = 1
        cur_set = set()

        # 题库中选题
        while(len(cur_set) < self.single_choice_count):
            question_num = random.randint(1, single_choice_sheet.max_row-2)
            # 该题号未被选过
            if question_num not in cur_set:
                cur_set.add(question_num)
                cur_question = str(cur_num) + u'、' + single_choice_sheet.cell(row=question_num + 2, column=4).value + u'\n'
                cur_question += u'   A、%s\n   B、%s\n   C、%s\n   D、%s\n' % (single_choice_sheet.cell(row=question_num + 2, column=5).value, single_choice_sheet.cell(row=question_num + 2, column=6).value,single_choice_sheet.cell(row=question_num + 2, column=7).value, single_choice_sheet.cell(row=question_num + 2, column=8).value)
                # 题目
                questions.add_paragraph(cur_question)
                # answer
                cur_answer += str(cur_num) + u'、' + single_choice_sheet.cell(row=question_num + 2, column=10).value + u'   '
                cur_num += 1
        answers.add_paragraph(cur_answer)
        
    def generateMultipleChoose(self, questions, answers, multiple_choice_set):
        cur_num = 1
        cur_answer = u''
        cur_set = set()

        # 题库中选题
        while(len(cur_set) < self.multiple_choice_count):
            question_num = random.randint(1, multiple_choice_set.max_row-2)
            # 该题号未被选过
            if question_num not in cur_set:
                cur_set.add(question_num)
                cur_question = str(cur_num) + u'、' + multiple_choice_set.cell(row=question_num + 2, column=4).value + u'\n'
                cur_question += u'   A、%s\n   B、%s\n   C、%s\n   D、%s\n' % (multiple_choice_set.cell(row=question_num + 2, column=5).value, multiple_choice_set.cell(row=question_num + 2, column=6).value,multiple_choice_set.cell(row=question_num + 2, column=7).value, multiple_choice_set.cell(row=question_num + 2, column=8).value)
                if multiple_choice_set.cell(row=question_num + 2, column=9).value:
                    cur_question += u'   E、' + multiple_choice_set.cell(row=question_num + 2, column=9).value + u'\n'
                
                questions.add_paragraph(cur_question)
                cur_answer += str(cur_num) + u'、' + multiple_choice_set.cell(row=question_num + 2, column=10).value + u'   '
                cur_num += 1

        answers.add_paragraph(cur_answer)
        pass

    def generateJudge(self, questions, answers, judge_sheet):
        cur_num = 1
        cur_set = set()
        cur_answer = u''
        while(len(cur_set) < self.judge_count):
            question_num = random.randint(1, judge_sheet.max_row-2)
            if question_num not in cur_set:
                cur_set.add(question_num)
                cur_question = str(cur_num) + u'、' + judge_sheet.cell(row=question_num + 2, column=4).value + u'（ ）'
                questions.add_paragraph(cur_question)
                cur_answer += str(cur_num) + u'、' + judge_sheet.cell(row=question_num + 2, column=10).value + u'   '
                cur_num += 1
        answers.add_paragraph(cur_answer)
        pass

    def generateQuestions(self):
        filepath = askopenfilename()
        
        # 题型
        workbook = load_workbook(filepath, data_only=True)
        sheets = workbook.get_sheet_names()
        single_choice_sheet = workbook.get_sheet_by_name(sheets[0])
        multiple_choice_set = workbook.get_sheet_by_name(sheets[1])
        judge_sheet = workbook.get_sheet_by_name(sheets[2])

        # worksheet， 单选
        rows = []
        for row in workbook.active.iter_rows():
            rows.append(row)

        # 生成self.count套试卷
        for count in range(1, self.count+1):
            # 当前试卷的问题word和答案的word
            questions = Document()
            answers = Document()
            
            # 设置标题
            head_paragraph = questions.add_paragraph('')
            head_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            head_paragraph.add_run(u'Examinations'+str(count)+u'\n')

            danxuanti_para = u'第一部分  单选题(20题)'
            questions.add_paragraph(danxuanti_para)
            answers.add_paragraph(danxuanti_para)
            self.generateSingleChoose(questions, answers, single_choice_sheet)
            
            duoxuanti_para = u'第二部分  多选题(10题)'
            questions.add_paragraph(duoxuanti_para)
            answers.add_paragraph(duoxuanti_para)  # answer
            self.generateMultipleChoose(questions, answers, multiple_choice_set)

            panduanti_para = u'第三部分  判断题(10题)'
            questions.add_paragraph(panduanti_para)
            answers.add_paragraph(panduanti_para)  # answer
            self.generateJudge(questions, answers, judge_sheet)
            
            # 保存文件，求出当前路径
            for i in range(len(filepath) - 1, -1, -1):
                if (filepath[i] == '/'):
                    break
            savepath = u""
            for j in range(0, i + 1):
                savepath += filepath[j]
            answer_savepath = savepath + u'Test' + str(count) + u'answer.docx'
            savepath = savepath + u'Test' + str(count) + u'.docx'
            questions.save(savepath)
            answers.save(answer_savepath)
            print("Finish generating!")
